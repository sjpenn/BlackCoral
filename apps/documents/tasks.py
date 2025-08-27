"""
Celery tasks for document processing.
Handles fetching, parsing, and extracting content from opportunity documents.
"""

import logging
import os
import requests
from typing import List, Dict, Any, Optional
from celery import shared_task
from django.core.files.base import ContentFile
from django.utils import timezone

from .models import Document
from apps.opportunities.models import Opportunity

logger = logging.getLogger('blackcoral.documents.tasks')


@shared_task(bind=True, max_retries=3)
def fetch_opportunity_documents(self, opportunity_id: int, document_links: Optional[List[Dict]] = None):
    """
    Fetch and store documents for an opportunity.
    
    Args:
        opportunity_id: ID of the opportunity
        document_links: Optional list of document links from API
    """
    try:
        opportunity = Opportunity.objects.get(id=opportunity_id)
        
        # Get document links from opportunity if not provided
        if not document_links:
            document_links = []
            
            # Extract from raw_data if available
            if opportunity.raw_data:
                # Resource links
                for link in opportunity.raw_data.get('resourceLinks', []):
                    document_links.append({
                        'url': link,
                        'type': 'resource',
                        'name': link.split('/')[-1] if '/' in link else 'Resource'
                    })
                
                # Additional info link
                if opportunity.raw_data.get('additionalInfoLink'):
                    document_links.append({
                        'url': opportunity.raw_data['additionalInfoLink'],
                        'type': 'additional_info',
                        'name': 'Additional Information'
                    })
        
        if not document_links:
            logger.info(f"No documents found for opportunity {opportunity_id}")
            opportunity.documents_fetched = True
            opportunity.save()
            return {'status': 'success', 'documents_count': 0}
        
        # Fetch each document
        fetched_count = 0
        for doc_info in document_links:
            try:
                # Check if document already exists
                if Document.objects.filter(
                    opportunity=opportunity,
                    source_url=doc_info['url']
                ).exists():
                    logger.debug(f"Document already exists: {doc_info['url']}")
                    continue
                
                # Fetch document
                fetch_single_document.delay(
                    opportunity_id=opportunity_id,
                    url=doc_info['url'],
                    doc_type=doc_info.get('type', 'unknown'),
                    doc_name=doc_info.get('name', 'Document')
                )
                fetched_count += 1
                
            except Exception as e:
                logger.error(f"Error queuing document fetch: {e}")
        
        opportunity.documents_fetched = True
        opportunity.save()
        
        return {
            'status': 'success',
            'documents_queued': fetched_count,
            'total_links': len(document_links)
        }
        
    except Opportunity.DoesNotExist:
        logger.error(f"Opportunity {opportunity_id} not found")
        return {'status': 'error', 'message': 'Opportunity not found'}
    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        raise self.retry(exc=e, countdown=300)


@shared_task(bind=True, max_retries=3)
def fetch_single_document(self, opportunity_id: int, url: str, doc_type: str, doc_name: str):
    """
    Fetch a single document from a URL.
    """
    try:
        opportunity = Opportunity.objects.get(id=opportunity_id)
        
        # Download document
        response = requests.get(url, timeout=30, headers={
            'User-Agent': 'BLACK CORAL Government Contracting System'
        })
        response.raise_for_status()
        
        # Determine file type from content-type or URL
        content_type = response.headers.get('Content-Type', '')
        if 'pdf' in content_type:
            file_type = 'PDF'
        elif 'word' in content_type or 'docx' in url:
            file_type = 'DOCX'
        elif 'excel' in content_type or 'xlsx' in url or 'xls' in url:
            file_type = 'XLS'
        elif 'zip' in content_type:
            file_type = 'ZIP'
        elif 'html' in content_type:
            file_type = 'HTML'
        else:
            file_type = 'OTHER'
        
        # Create document record
        document = Document.objects.create(
            opportunity=opportunity,
            title=doc_name,
            file_type=file_type,
            file_size=len(response.content),
            source_url=url,
            document_type=doc_type,
            processing_status='pending'
        )
        
        # Save file
        file_name = f"{opportunity.solicitation_number}_{document.id}.{file_type.lower()}"
        document.file_path.save(file_name, ContentFile(response.content))
        
        logger.info(f"Downloaded document: {file_name}")
        
        # Queue for processing
        parse_document.delay(document.id)
        
        return {
            'status': 'success',
            'document_id': document.id,
            'file_type': file_type
        }
        
    except requests.RequestException as e:
        logger.error(f"Error downloading document from {url}: {e}")
        raise self.retry(exc=e, countdown=600)
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise self.retry(exc=e, countdown=300)


@shared_task(bind=True, max_retries=3)
def parse_document(self, document_id: int):
    """
    Parse and extract text from a document.
    """
    try:
        document = Document.objects.get(id=document_id)
        document.processing_status = 'processing'
        document.save()
        
        # Route to appropriate parser based on file type
        if document.file_type == 'PDF':
            text = parse_pdf_document(document)
        elif document.file_type in ['DOCX', 'DOC']:
            text = parse_word_document(document)
        elif document.file_type == 'HTML':
            text = parse_html_document(document)
        elif document.file_type in ['XLS', 'XLSX']:
            text = parse_excel_document(document)
        else:
            text = ''
            logger.warning(f"Unsupported file type: {document.file_type}")
        
        # Save extracted text
        document.extracted_text = text
        document.processing_status = 'completed'
        document.processed_at = timezone.now()
        document.save()
        
        logger.info(f"Parsed document {document_id}: {len(text)} characters extracted")
        
        # Trigger AI analysis if text was extracted
        if text and len(text) > 100:
            # Queue chunking and embedding generation
            create_document_chunks.delay(document_id)
            # Queue comprehensive analysis
            analyze_document_with_ai.delay(document_id, 'comprehensive')
        
        return {
            'status': 'success',
            'document_id': document_id,
            'text_length': len(text)
        }
        
    except Document.DoesNotExist:
        logger.error(f"Document {document_id} not found")
        return {'status': 'error', 'message': 'Document not found'}
    except Exception as e:
        logger.error(f"Error parsing document: {e}")
        document.processing_status = 'failed'
        document.processing_error = str(e)
        document.save()
        raise self.retry(exc=e, countdown=600)


def parse_pdf_document(document: Document) -> str:
    """Parse PDF document and extract text."""
    try:
        import fitz  # PyMuPDF
        
        text = []
        with fitz.open(document.file_path.path) as pdf:
            for page_num, page in enumerate(pdf):
                page_text = page.get_text()
                if page_text:
                    text.append(f"Page {page_num + 1}:\n{page_text}")
        
        return '\n\n'.join(text)
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        return ''


def parse_word_document(document: Document) -> str:
    """Parse Word document and extract text."""
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(document.file_path.path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                if row_text.strip():
                    text.append(row_text)
        
        return '\n\n'.join(text)
    except Exception as e:
        logger.error(f"Error parsing Word document: {e}")
        return ''


def parse_html_document(document: Document) -> str:
    """Parse HTML document and extract text."""
    try:
        from bs4 import BeautifulSoup
        
        with open(document.file_path.path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text and clean it up
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    except Exception as e:
        logger.error(f"Error parsing HTML: {e}")
        return ''


def parse_excel_document(document: Document) -> str:
    """Parse Excel document and extract text."""
    try:
        import pandas as pd
        
        # Read all sheets
        excel_file = pd.ExcelFile(document.file_path.path)
        text = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name)
            text.append(f"Sheet: {sheet_name}")
            text.append(df.to_string())
        
        return '\n\n'.join(text)
    except Exception as e:
        logger.error(f"Error parsing Excel: {e}")
        return ''


@shared_task
def process_pending_documents():
    """
    Process all pending documents (runs periodically).
    """
    pending_docs = Document.objects.filter(processing_status='pending')
    
    for doc in pending_docs[:10]:  # Process max 10 at a time
        parse_document.delay(doc.id)
    
    return {
        'status': 'success',
        'documents_queued': min(pending_docs.count(), 10)
    }


@shared_task
def cleanup_old_documents():
    """
    Clean up old documents from inactive opportunities.
    """
    from datetime import timedelta
    
    cutoff_date = timezone.now() - timedelta(days=180)
    
    old_documents = Document.objects.filter(
        opportunity__is_active=False,
        created_at__lt=cutoff_date
    )
    
    count = 0
    for doc in old_documents:
        if doc.file_path:
            doc.file_path.delete()
        doc.delete()
        count += 1
    
    return {
        'status': 'success',
        'deleted_count': count
    }


@shared_task(bind=True, max_retries=3)
def process_uploaded_document(self, upload_id: int, file_content_hex: str):
    """
    Process an uploaded document with security validation and parsing.
    
    Args:
        upload_id: DocumentUpload record ID
        file_content_hex: File content as hex string
    """
    try:
        from .models import DocumentUpload, DocumentSecurity
        from .security import FileValidationService, VirusScanService
        
        upload = DocumentUpload.objects.get(id=upload_id)
        upload.status = 'validating'
        upload.save()
        
        # Convert hex back to bytes
        file_content = bytes.fromhex(file_content_hex)
        
        # File validation
        validator = FileValidationService()
        validation_result = validator.validate_file(
            content=file_content,
            filename=upload.original_filename,
            mime_type=upload.mime_type
        )
        
        if not validation_result['is_valid']:
            upload.status = 'rejected'
            upload.validation_errors = validation_result['errors']
            upload.error_message = 'File validation failed'
            upload.save()
            return {'status': 'rejected', 'errors': validation_result['errors']}
        
        # Virus scanning
        upload.status = 'scanning'
        upload.save()
        
        scanner = VirusScanService()
        scan_result = scanner.scan_content(file_content)
        
        upload.virus_scan_result = scan_result
        upload.is_safe = scan_result.get('is_safe', False)
        
        if not upload.is_safe:
            upload.status = 'rejected'
            upload.error_message = 'File failed virus scan'
            upload.security_flags = scan_result.get('threats', [])
            upload.save()
            return {'status': 'rejected', 'threats': scan_result.get('threats', [])}
        
        # Create Document record
        upload.status = 'completed'
        upload.upload_progress = 100
        
        document = Document.objects.create(
            opportunity=upload.opportunity,
            title=upload.original_filename,
            file_type=upload.mime_type.split('/')[-1].upper(),
            file_size=upload.file_size,
            document_type='uploaded',
            processing_status='pending'
        )
        
        # Save file
        from django.core.files.base import ContentFile
        file_name = f"upload_{upload.id}_{document.id}.{document.file_type.lower()}"
        document.file_path.save(file_name, ContentFile(file_content))
        
        # Create security record
        DocumentSecurity.objects.create(
            document=document,
            classification='unclassified',  # Default
            classification_authority=upload.user
        )
        
        upload.created_document = document
        upload.save()
        
        # Queue for parsing and analysis
        parse_document.delay(document.id)
        
        return {
            'status': 'success',
            'upload_id': upload.id,
            'document_id': document.id
        }
        
    except DocumentUpload.DoesNotExist:
        logger.error(f"DocumentUpload {upload_id} not found")
        return {'status': 'error', 'message': 'Upload not found'}
    except Exception as e:
        logger.error(f"Error processing upload {upload_id}: {e}")
        upload.status = 'failed'
        upload.error_message = str(e)
        upload.save()
        raise self.retry(exc=e, countdown=300)


@shared_task(bind=True, max_retries=3)
def generate_document_embeddings(self, document_id: int):
    """
    Generate vector embeddings for document chunks.
    """
    try:
        from .models import DocumentChunk
        from .vector_store.embeddings import EmbeddingService
        from .vector_store.chroma_store import ChromaVectorStore
        
        document = Document.objects.get(id=document_id)
        
        if not document.is_processed or not document.extracted_text:
            logger.warning(f"Document {document_id} not ready for embedding generation")
            return {'status': 'skipped', 'reason': 'Document not processed'}
        
        # Initialize services
        embedding_service = EmbeddingService()
        vector_store = ChromaVectorStore()
        
        # Get or create chunks
        chunks = document.chunks.all()
        if not chunks.exists():
            # Create chunks if they don't exist
            create_document_chunks.delay(document_id)
            return {'status': 'deferred', 'reason': 'Creating chunks first'}
        
        # Generate embeddings for each chunk
        processed_count = 0
        failed_count = 0
        
        for chunk in chunks:
            try:
                # Generate embedding
                embedding = embedding_service.generate_text_embedding(chunk.text)
                
                if embedding:
                    # Store in vector database
                    vector_store.add_document(
                        document_id=str(document.id),
                        chunk_id=str(chunk.id),
                        content=chunk.text,
                        metadata={
                            'document_title': document.title,
                            'page_number': chunk.page_number,
                            'section_title': chunk.section_title,
                            'chunk_index': chunk.chunk_index,
                            **chunk.metadata
                        },
                        embedding=embedding
                    )
                    
                    # Store embedding in chunk record
                    chunk.embedding_vector = embedding
                    chunk.embedding_model = embedding_service.model_name
                    chunk.save()
                    
                    processed_count += 1
                else:
                    failed_count += 1
                    logger.warning(f"Failed to generate embedding for chunk {chunk.id}")
                    
            except Exception as e:
                logger.error(f"Error processing chunk {chunk.id}: {e}")
                failed_count += 1
        
        # Update document embedding status
        document.has_embeddings = processed_count > 0
        document.embedding_updated_at = timezone.now()
        document.save()
        
        return {
            'status': 'success',
            'document_id': document_id,
            'processed_chunks': processed_count,
            'failed_chunks': failed_count
        }
        
    except Document.DoesNotExist:
        logger.error(f"Document {document_id} not found")
        return {'status': 'error', 'message': 'Document not found'}
    except Exception as e:
        logger.error(f"Error generating embeddings: {e}")
        raise self.retry(exc=e, countdown=600)


@shared_task(bind=True, max_retries=3)
def create_document_chunks(self, document_id: int, chunk_size: int = 1000, overlap: int = 200):
    """
    Create text chunks from document content for RAG.
    """
    try:
        from .models import DocumentChunk
        from .vector_store.chunking import TextChunker
        
        document = Document.objects.get(id=document_id)
        
        if not document.extracted_text:
            logger.warning(f"No extracted text for document {document_id}")
            return {'status': 'skipped', 'reason': 'No extracted text'}
        
        # Clear existing chunks
        document.chunks.all().delete()
        
        # Initialize chunker
        chunker = TextChunker(
            chunk_size=chunk_size,
            overlap=overlap,
            respect_sentence_boundaries=True
        )
        
        # Create chunks
        chunks = chunker.chunk_text(document.extracted_text)
        
        chunk_objects = []
        for i, chunk_text in enumerate(chunks):
            # Extract metadata from chunk position
            chunk_metadata = {
                'chunk_size': len(chunk_text),
                'word_count': len(chunk_text.split()),
                'char_start': sum(len(c) for c in chunks[:i]),
                'char_end': sum(len(c) for c in chunks[:i+1])
            }
            
            chunk_obj = DocumentChunk(
                document=document,
                chunk_index=i,
                text=chunk_text,
                metadata=chunk_metadata
            )
            chunk_objects.append(chunk_obj)
        
        # Bulk create chunks
        DocumentChunk.objects.bulk_create(chunk_objects)
        
        # Queue embedding generation
        generate_document_embeddings.delay(document_id)
        
        return {
            'status': 'success',
            'document_id': document_id,
            'chunks_created': len(chunk_objects)
        }
        
    except Document.DoesNotExist:
        logger.error(f"Document {document_id} not found")
        return {'status': 'error', 'message': 'Document not found'}
    except Exception as e:
        logger.error(f"Error creating chunks: {e}")
        raise self.retry(exc=e, countdown=300)


@shared_task(bind=True, max_retries=3)
def analyze_document_with_ai(self, document_id: int, analysis_type: str = 'comprehensive', 
                           parameters: Optional[Dict] = None, priority: str = 'normal'):
    """
    Perform AI-powered analysis of a document.
    """
    try:
        from .models import DocumentAnalysisResult, DocumentExtractionSession
        from apps.ai_integration.services import AIAnalysisService
        
        document = Document.objects.get(id=document_id)
        
        if not document.extracted_text:
            logger.warning(f"No extracted text for document {document_id}")
            return {'status': 'skipped', 'reason': 'No extracted text'}
        
        # Initialize AI service
        ai_service = AIAnalysisService()
        
        # Create extraction session
        session = DocumentExtractionSession.objects.create(
            document=document,
            extraction_type=analysis_type,
            status='processing',
            started_at=timezone.now(),
            extraction_parameters=parameters or {},
            ai_provider=ai_service.provider_name
        )
        
        try:
            # Perform analysis based on type
            if analysis_type == 'comprehensive':
                result = ai_service.comprehensive_analysis(
                    text=document.extracted_text,
                    document_title=document.title,
                    parameters=parameters
                )
            elif analysis_type == 'requirements_extraction':
                result = ai_service.extract_requirements(
                    text=document.extracted_text,
                    parameters=parameters
                )
            elif analysis_type == 'compliance_check':
                result = ai_service.compliance_analysis(
                    text=document.extracted_text,
                    parameters=parameters
                )
            elif analysis_type == 'technical_specifications':
                result = ai_service.extract_technical_specs(
                    text=document.extracted_text,
                    parameters=parameters
                )
            else:
                raise ValueError(f"Unknown analysis type: {analysis_type}")
            
            # Update session with results
            session.status = 'completed'
            session.completed_at = timezone.now()
            session.processing_time = session.completed_at - session.started_at
            session.extracted_data = result.get('data', {})
            session.confidence_scores = result.get('confidence_scores', {})
            session.save()
            
            # Create or update analysis result record
            analysis_result, created = DocumentAnalysisResult.objects.update_or_create(
                document=document,
                analysis_type=analysis_type,
                defaults={
                    'results': result.get('data', {}),
                    'confidence_score': result.get('overall_confidence', 0.0),
                    'processing_time': session.processing_time,
                    'ai_provider': ai_service.provider_name,
                    'model_version': ai_service.model_version,
                    'analysis_parameters': parameters or {},
                    'accuracy_score': result.get('accuracy_score'),
                    'completeness_score': result.get('completeness_score')
                }
            )
            
            # Update document fields based on analysis
            if analysis_type == 'comprehensive':
                document.summary = result.get('data', {}).get('summary', '')
                document.key_requirements = result.get('data', {}).get('requirements', [])
                document.compliance_notes = result.get('data', {}).get('compliance_notes', [])
                document.ai_analysis_results = result.get('data', {})
                
                # Update boolean flags
                document.is_sow = result.get('data', {}).get('is_sow', False)
                document.is_pws = result.get('data', {}).get('is_pws', False)
                document.contains_requirements = result.get('data', {}).get('has_requirements', False)
                document.contains_technical_specs = result.get('data', {}).get('has_technical_specs', False)
                document.contains_terms_conditions = result.get('data', {}).get('has_terms_conditions', False)
                document.contains_pricing_info = result.get('data', {}).get('has_pricing_info', False)
                
                document.save()
            
            return {
                'status': 'success',
                'document_id': document_id,
                'analysis_type': analysis_type,
                'session_id': session.id,
                'result_id': analysis_result.id,
                'confidence_score': result.get('overall_confidence', 0.0)
            }
            
        except Exception as e:
            # Update session with error
            session.status = 'failed'
            session.completed_at = timezone.now()
            session.errors = [str(e)]
            session.save()
            raise e
        
    except Document.DoesNotExist:
        logger.error(f"Document {document_id} not found")
        return {'status': 'error', 'message': 'Document not found'}
    except Exception as e:
        logger.error(f"Error analyzing document: {e}")
        raise self.retry(exc=e, countdown=600)


@shared_task(bind=True, max_retries=3)
def create_processing_pipeline(self, document_id: int, pipeline_type: str = 'full_analysis',
                              priority: str = 'normal', user_id: Optional[int] = None):
    """
    Create and execute a document processing pipeline.
    """
    try:
        from .models import DocumentProcessingPipeline
        
        document = Document.objects.get(id=document_id)
        
        # Define pipeline stages based on type
        if pipeline_type == 'full_analysis':
            stages = [
                {'name': 'text_extraction', 'task': 'parse_document'},
                {'name': 'chunking', 'task': 'create_document_chunks'},
                {'name': 'embedding_generation', 'task': 'generate_document_embeddings'},
                {'name': 'comprehensive_analysis', 'task': 'analyze_document_with_ai', 'params': {'analysis_type': 'comprehensive'}},
                {'name': 'requirements_extraction', 'task': 'analyze_document_with_ai', 'params': {'analysis_type': 'requirements_extraction'}},
                {'name': 'compliance_check', 'task': 'analyze_document_with_ai', 'params': {'analysis_type': 'compliance_check'}}
            ]
        elif pipeline_type == 'basic_processing':
            stages = [
                {'name': 'text_extraction', 'task': 'parse_document'},
                {'name': 'chunking', 'task': 'create_document_chunks'},
                {'name': 'embedding_generation', 'task': 'generate_document_embeddings'}
            ]
        elif pipeline_type == 'analysis_only':
            stages = [
                {'name': 'comprehensive_analysis', 'task': 'analyze_document_with_ai', 'params': {'analysis_type': 'comprehensive'}}
            ]
        else:
            raise ValueError(f"Unknown pipeline type: {pipeline_type}")
        
        # Create pipeline record
        pipeline = DocumentProcessingPipeline.objects.create(
            document=document,
            pipeline_name=pipeline_type,
            status='running',
            stages=stages,
            total_stages=len(stages),
            started_at=timezone.now()
        )
        
        # Execute stages sequentially
        for i, stage in enumerate(stages):
            try:
                pipeline.current_stage = i
                pipeline.save()
                
                # Execute stage task
                task_name = stage['task']
                task_params = stage.get('params', {})
                
                if task_name == 'parse_document':
                    result = parse_document.delay(document_id)
                elif task_name == 'create_document_chunks':
                    result = create_document_chunks.delay(document_id)
                elif task_name == 'generate_document_embeddings':
                    result = generate_document_embeddings.delay(document_id)
                elif task_name == 'analyze_document_with_ai':
                    result = analyze_document_with_ai.delay(document_id, **task_params)
                else:
                    raise ValueError(f"Unknown task: {task_name}")
                
                # Wait for task completion (with timeout)
                task_result = result.get(timeout=300)  # 5 minute timeout per stage
                
                # Store stage result
                pipeline.stage_results[stage['name']] = task_result
                pipeline.stages_completed = i + 1
                pipeline.progress_percentage = int((i + 1) / len(stages) * 100)
                pipeline.save()
                
                if task_result.get('status') != 'success':
                    raise Exception(f"Stage {stage['name']} failed: {task_result}")
                
            except Exception as e:
                # Log error and mark pipeline as failed
                error_entry = {
                    'stage': stage['name'],
                    'error': str(e),
                    'timestamp': timezone.now().isoformat()
                }
                
                pipeline.error_log.append(error_entry)
                pipeline.status = 'failed'
                pipeline.completed_at = timezone.now()
                pipeline.save()
                
                raise e
        
        # Mark pipeline as completed
        pipeline.status = 'completed'
        pipeline.completed_at = timezone.now()
        pipeline.progress_percentage = 100
        pipeline.save()
        
        return {
            'status': 'success',
            'pipeline_id': pipeline.id,
            'document_id': document_id,
            'stages_completed': len(stages)
        }
        
    except Document.DoesNotExist:
        logger.error(f"Document {document_id} not found")
        return {'status': 'error', 'message': 'Document not found'}
    except Exception as e:
        logger.error(f"Error in processing pipeline: {e}")
        raise self.retry(exc=e, countdown=600)